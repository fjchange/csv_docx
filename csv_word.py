import docx
import csv
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

path_pre='H:\\Users\\kiwi feng\\Downloads\\计算机科学与工程学院本科生导师制、毕业实习、毕业设计答辩工作方案 (2)\\'
docx_path='本科毕业实习日志-专业实习.docx'
csv_path='专业实习.csv'
docx_save_path='日志\\本科毕业实习日志-专业实习'
def main():
    with open(path_pre+csv_path) as c:
        lines=csv.reader(c)
        next_day_to_do=''
        for i,line in enumerate(lines):
            j=1
            document = docx.Document(path_pre + docx_path)
            document.tables[0].cell(4,1).paragraphs[0].add_run(line[0])
            cell_now=document.tables[0].cell(5,1).paragraphs[0]
            cell_now.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
            cell_now.add_run('\n')
            style_now=cell_now.runs[0].style
            font_size=cell_now.runs[0].font.size
            if i ==0:
                cell_now.add_run('1. 阅读Faster-RCNN论文 \n 2. 阅读Faster-RCNN实现源码\n')
            else:
                cell_now.add_run(next_day_to_do+'\n')

            cell_now.add_run('二、今日工作进展：\n',style=style_now)
            cell_now.runs[-1].bold=True
            cell_now.runs[-1].font.size=font_size

            cell_now.add_run(line[2]+'\n',style_now)

            cell_now.add_run('三、明日工作计划：\n',style_now)
            cell_now.runs[-1].bold=True
            cell_now.runs[-1].font.size=font_size
            cell_now.add_run(line[3]+'\n')
            next_day_to_do=line[3]
            cell_now.add_run('四、其他（存在问题、体会、提议等）:\n',style_now)
            cell_now.runs[-1].bold=True
            cell_now.runs[-1].font.size=font_size
            if len(line)<5:
                cell_now.add_run('无\n')
            else:
                cell_now.add_run(line[-1]+'\n')

            document.save(path_pre+docx_save_path+'-'+line[0]+'.docx')
            print(i,'successfully done!------')

if __name__=='__main__':
    main()